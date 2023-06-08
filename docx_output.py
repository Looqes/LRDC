import docx
from docx.enum.text import WD_COLOR_INDEX

colors = [WD_COLOR_INDEX.BRIGHT_GREEN, WD_COLOR_INDEX.VIOLET, WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.YELLOW]

def getcolors(colors):
    return (color for color in colors)

def write_output_docx(expression1, expression2, clause_overlap, clause_non_overlap, result, possible_clause_differences):
    # Create an instance of a word document
    doc = docx.Document()
    
    # Add a Title to the document 
    doc.add_heading('Expression difference', 2)
    doc.add_paragraph().add_run(str(expression1))
    doc.add_paragraph().add_run(str(expression2))

    doc.add_paragraph()
    doc.add_paragraph().add_run("Overlap between rulesets: ")
    doc.add_paragraph().add_run(str(clause_overlap))
    doc.add_paragraph().add_run(str(clause_overlap)).add_break()


    doc.add_paragraph().add_run("Changed rules: ")
    for matching in result:
        # print(possible_clause_differences[matching])
        clause_difference = possible_clause_differences[matching]
        overlap = list(possible_clause_differences[matching].overlap)
        negations = list(possible_clause_differences[matching].negations)
        deletions = list(possible_clause_differences[matching].deletions)
        additions = list(possible_clause_differences[matching].additions)

        # print(overlap)

        first_rule = str(possible_clause_differences[matching].clause)
        first_rule_length = possible_clause_differences[matching].clause.length

        second_rule = str(possible_clause_differences[matching].target)
        second_rule_length = possible_clause_differences[matching].clause.length

        visualization = doc.add_paragraph()

        # first clause
        visualization.add_run("(")
        highlight_colors = getcolors(colors)
        i = 1
        for literal in overlap:
            visualization.add_run(str(literal))

            if i < first_rule_length:
                visualization.add_run(" ∨ ")
                i += 1

        for negated_pair in negations:
            visualization.add_run(str(negated_pair[0])).font.highlight_color = next(highlight_colors)

            if i < first_rule_length:
                visualization.add_run(" ∨ ")
                i += 1

        for addition in additions:
            visualization.add_run(str(addition))

            if i < first_rule_length:
                visualization.add_run(" ∨ ")
                i += 1    

        visualization.add_run(")").add_break()


        # second clause
        visualization.add_run("(")
        highlight_colors = getcolors(colors)
        i = 1
        for literal in overlap:
            visualization.add_run(str(literal))

            if i < second_rule_length:
                visualization.add_run(" ∨ ")
                i += 1

        for negated_pair in negations:
            visualization.add_run(str(negated_pair[1])).font.highlight_color = next(highlight_colors)

            if i < second_rule_length:
                visualization.add_run(" ∨ ")
                i += 1

        for deletion in deletions:
            visualization.add_run(str(deletion))

            if i < second_rule_length:
                visualization.add_run(" ∨ ")
                i += 1    

        visualization.add_run(")")
            
        visualization.add_run().add_break()

    deleted_clauses = doc.add_paragraph()
    deleted_clauses.add_run("Rules that no longer appear in the second ruleset").add_break()
    
    if clause_non_overlap[0].clauses:
        deleted_clauses.add_run(str(clause_non_overlap[0])).add_break()
    else:
        deleted_clauses.add_run("    None...").add_break()

    new_clauses = doc.add_paragraph()
    new_clauses.add_run("New Rules in the second ruleset").add_break()
    if clause_non_overlap[1].clauses:
        new_clauses.add_run(str(clause_non_overlap[1]))
    else:
        new_clauses.add_run("    No new rules...")




    # Creating paragraph with some content and Highlighting it.
    # highlight_para = doc.add_paragraph(
    #        ).add_run(
    #            str(expression1)
    #                  ).font.highlight_color = next(highlight_colors)
    
    # Now save the document to a location 
    doc.save('gfg.docx')




